# services/file_service.py
from __future__ import annotations
import os, io, math, json, subprocess, shutil, tempfile
from pathlib import Path
from typing import Dict, Any, List, Optional
import fitz
from PIL import Image

from html2text import html2text

# DOCX support
try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn as docx_qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# XLSX support
try:
    from openpyxl import load_workbook
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

# 统一的根目录：每个 fileId 一个子目录
DATA_ROOT = Path("data")

def workdir(file_id: str) -> Path:
    d = DATA_ROOT / file_id
    d.mkdir(parents=True, exist_ok=True)
    return d

def dir_original_pages(file_id: str) -> Path:
    p = workdir(file_id) / "pages" / "original"
    p.mkdir(parents=True, exist_ok=True)
    return p

def dir_parsed_pages(file_id: str) -> Path:
    p = workdir(file_id) / "pages" / "parsed"
    p.mkdir(parents=True, exist_ok=True)
    return p

def get_actual_file_path(file_id: str, ext: str) -> Path:
    """获取实际上传的文件路径，支持新旧命名方式"""
    new_path = workdir(file_id) / ("original" + ext)
    if new_path.exists():
        return new_path
    old_path = workdir(file_id) / "original.pdf"
    if old_path.exists():
        return old_path
    return new_path

def find_actual_file(file_id: str) -> Path:
    """查找实际上传的文件（任意扩展名）"""
    wd = workdir(file_id)
    for ext in [".pdf", ".docx", ".xlsx", ".txt", ".md"]:
        p = wd / ("original" + ext)
        if p.exists():
            return p
    old = wd / "original.pdf"
    if old.exists():
        return old
    return wd / "original.pdf"

def detect_file_type(filename: str) -> str:
    """根据文件扩展名检测文件类型"""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    elif ext == ".docx":
        return "docx"
    elif ext == ".xlsx":
        return "xlsx"
    elif ext == ".txt":
        return "txt"
    elif ext == ".md":
        return "md"
    return "pdf"

def save_upload(file_id: str, upload_bytes: bytes, filename: str) -> Dict[str, Any]:
    """保存上传文件，并根据类型返回元信息"""
    file_type = detect_file_type(filename)
    ext = Path(filename).suffix.lower()
    file_path = workdir(file_id) / ("original" + ext)
    file_path.write_bytes(upload_bytes)

    pages = 0
    if file_type == "pdf":
        with fitz.open(file_path) as doc:
            pages = doc.page_count

    return {"fileId": file_id, "name": filename, "pages": pages, "fileType": file_type}

def markdown_output(file_id: str) -> Path:
    return workdir(file_id) / "output.md"

def images_dir(file_id: str) -> Path:
    p = workdir(file_id) / "images"
    p.mkdir(parents=True, exist_ok=True)
    return p

def parse_pdf_with_mineru(file_id: str) -> Dict[str, Any]:
    """使用 MinerU 解析 PDF 文件"""
    pdf_path = get_actual_file_path(file_id, ".pdf")
    output_dir = workdir(file_id)
    
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_output_path = Path(temp_dir)
        
        cmd = [
            "mineru",
            "-p", str(pdf_path),
            "-o", str(temp_output_path),
            "-b", "pipeline"
        ]
        
        try:
            print(f"执行 MinerU 命令: {' '.join(cmd)}")
            result = subprocess.run(
                cmd,
                check=True,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace"
            )
            print(f"MinerU 输出 (stdout): {result.stdout}")
            if result.stderr:
                print(f"MinerU 输出 (stderr): {result.stderr}")
        except subprocess.CalledProcessError as e:
            print(f"MinerU 解析失败: {e.stderr}")
            raise RuntimeError(f"PDF 解析失败: {e.stderr}") from e
        
        print(f"临时目录内容: {list(temp_output_path.iterdir())}")
        
        # 查找 .md 文件，包括在子目录中
        md_files = list(temp_output_path.rglob("*.md"))
        print(f"找到 Markdown 文件: {md_files}")
        
        if not md_files:
            print(f"临时目录完整结构:")
            for item in temp_output_path.rglob("*"):
                print(f"  {item}")
            raise RuntimeError("MinerU 未生成 Markdown 文件")
        
        md_file = md_files[0]
        target_md = markdown_output(file_id)
        print(f"复制 Markdown: {md_file} -> {target_md}")
        shutil.copy2(md_file, target_md)
        
        for item in temp_output_path.iterdir():
            if item.is_dir():
                auto_dir = item / "auto"
                img_src = None
                if auto_dir.exists() and (auto_dir / "images").exists():
                    img_src = auto_dir / "images"
                elif (item / "images").exists():
                    img_src = item / "images"
                
                if img_src:
                    for img in img_src.iterdir():
                        if img.is_file():
                            shutil.copy2(img, images_dir(file_id) / img.name)
    
    return {"markdown": str(target_md), "images_dir": "images"}

def render_original_pages(file_id: str, dpi: int = 144):
    """把原始 PDF 渲染为 PNG，存到 pages/original/"""
    pdf_path = get_actual_file_path(file_id, ".pdf")
    out_dir = dir_original_pages(file_id)
    with fitz.open(pdf_path) as doc:
        for idx, page in enumerate(doc, start=1):
            mat = fitz.Matrix(dpi/72, dpi/72)
            pix = page.get_pixmap(matrix=mat)
            (out_dir / f"page-{idx:04d}.png").write_bytes(pix.tobytes("png"))

# ---- DOCX / TXT / MD parsing helpers ----

def _get_heading_prefix(level: int) -> str:
    level = max(1, min(level, 6))
    return "#" * level + " "

def _table_to_markdown(table) -> str:
    rows = table.rows
    if not rows:
        return ""
    md_rows = []
    header_cells = rows[0].cells
    header = "| " + " | ".join(cell.text.replace("|", "\\|").replace("\n", " ") for cell in header_cells) + " |"
    md_rows.append(header)
    separator = "| " + " | ".join("---" for _ in header_cells) + " |"
    md_rows.append(separator)
    for row in rows[1:]:
        cells = row.cells
        row_text = "| " + " | ".join(cell.text.replace("|", "\\|").replace("\n", " ") for cell in cells) + " |"
        md_rows.append(row_text)
    return "\n".join(md_rows)

def parse_docx(file_id: str) -> Dict[str, Any]:
    if not HAS_DOCX:
        raise ImportError("python-docx is not installed")

    file_path = get_actual_file_path(file_id, ".docx")
    doc = DocxDocument(str(file_path))
    md_lines: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue

        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", ""))
                md_lines.append(_get_heading_prefix(level) + text + "\n")
            except ValueError:
                md_lines.append("## " + text + "\n")
        else:
            has_bold = any(run.bold for run in para.runs if run.text.strip())
            if has_bold:
                md_lines.append("**" + text + "**\n")
            else:
                md_lines.append(text + "\n")

    for table in doc.tables:
        md_table = _table_to_markdown(table)
        if md_table:
            md_lines.append(md_table + "\n")

    img_dir = images_dir(file_id)
    img_shape_map = {}
    img_counter = 0

    for para_idx, para in enumerate(doc.paragraphs):
        for run in para.runs:
            drawings = run._element.findall('.//a:blip', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
            for blip in drawings:
                rel_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rel_id and rel_id in doc.part.rels:
                    image_part = doc.part.rels[rel_id].target_part
                    img_bytes = image_part.blob
                    img_filename = f"docx_img_{img_counter}.png"
                    img_path = img_dir / img_filename
                    img = Image.open(io.BytesIO(img_bytes))
                    if img.mode in ('RGBA', 'P'):
                        img = img.convert('RGB')
                    img.save(str(img_path), 'PNG')
                    img_shape_map.setdefault(para_idx, []).append(img_filename)
                    img_counter += 1

    if img_shape_map:
        new_lines: List[str] = []
        para_count = len(doc.paragraphs)
        for i in range(para_count):
            new_lines.append(md_lines[i])
            if i in img_shape_map:
                for img_filename in img_shape_map[i]:
                    new_lines.append(f"![Image](./images/{img_filename})\n")
        new_lines.extend(md_lines[para_count:])
        md_lines = new_lines

    out_md = markdown_output(file_id)
    out_md.write_text("\n".join(md_lines), encoding="utf-8")
    return {"markdown": out_md.name, "images_dir": "images"}

def parse_text_file(file_id: str, file_type: str = "txt") -> Dict[str, Any]:
    ext = f".{file_type}"
    file_path = get_actual_file_path(file_id, ext)

    raw_bytes = file_path.read_bytes()
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            content = raw_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        content = raw_bytes.decode("utf-8", errors="replace")

    if file_type == "md":
        md_content = content
    else:
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        md_content = "\n\n".join(paragraphs)

    out_md = markdown_output(file_id)
    out_md.write_text(md_content, encoding="utf-8")
    return {"markdown": out_md.name, "images_dir": "images"}

def parse_xlsx(file_id: str) -> Dict[str, Any]:
    if not HAS_XLSX:
        raise ImportError("openpyxl is not installed")

    file_path = get_actual_file_path(file_id, ".xlsx")
    wb = load_workbook(filename=str(file_path), read_only=True, data_only=True)
    ws = wb.active
    
    md_lines: List[str] = []
    rows = list(ws.iter_rows(values_only=False))
    
    if not rows:
        wb.close()
        out_md = markdown_output(file_id)
        out_md.write_text("", encoding="utf-8")
        return {"markdown": out_md.name, "images_dir": "images"}
    
    header_row = rows[0]
    headers = [cell.value for cell in header_row]
    
    for row in rows[1:]:
        values = [cell.value for cell in row]
        if not any(v is not None and str(v).strip() for v in values):
            continue
        
        parts = []
        for h, v in zip(headers, values):
            h_str = str(h).strip() if h else ""
            v_str = str(v).strip() if v is not None else ""
            if h_str:
                parts.append(f"{h_str}：{v_str}")
        
        if parts:
            line = "；".join(parts)
            md_lines.append(line + "\n")
    
    wb.close()
    
    out_md = markdown_output(file_id)
    out_md.write_text("\n".join(md_lines), encoding="utf-8")
    return {"markdown": out_md.name, "images_dir": "images"}

def run_full_parse_pipeline(file_id: str) -> Dict[str, Any]:
    """
    完整解析流程：根据文件类型自动选择解析方式
    """
    file_path = find_actual_file(file_id)
    file_type = detect_file_type(file_path.name)

    if file_type == "pdf":
        render_original_pages(file_id)
        md_info = parse_pdf_with_mineru(file_id)
    elif file_type == "docx":
        md_info = parse_docx(file_id)
    elif file_type == "xlsx":
        md_info = parse_xlsx(file_id)
    elif file_type in ("txt", "md"):
        md_info = parse_text_file(file_id, file_type)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    return {"md": md_info["markdown"]}